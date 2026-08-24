"""Verify the real DRAFT_REVIEW API -> Worker -> parser -> result loop."""

from __future__ import annotations

import os
import tempfile
import threading
import time
from functools import partial
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

import httpx
from docx import Document

BASE_URL = os.getenv("SMOKE_BASE_URL", "http://127.0.0.1:8000")
FIXTURE_HOST = os.getenv("SMOKE_FIXTURE_HOST", "api")
FIXTURE_PORT = int(os.getenv("SMOKE_FIXTURE_PORT", "18080"))


class QuietHandler(SimpleHTTPRequestHandler):
    def log_message(self, format: str, *args: object) -> None:
        return None


def write_docx(path: Path, title: str) -> None:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("这是完全合成的工程冒烟测试内容。")
    document.save(path)


def payload() -> dict:
    base = f"http://{FIXTURE_HOST}:{FIXTURE_PORT}"
    return {
        "client_reference_id": "docker-smoke-draft-real-parse",
        "target_file": {"url": f"{base}/target.docx", "file_name": "target.docx"},
        "template_file": {"url": f"{base}/template.docx", "file_name": "template.docx"},
        "reference_files": [
            {"url": f"{base}/reference.docx", "file_name": "reference.docx"}
        ],
    }


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="contract-review-smoke-") as directory:
        root = Path(directory)
        write_docx(root / "target.docx", "融资租赁合同")
        write_docx(root / "template.docx", "融资租赁合同")
        write_docx(root / "reference.docx", "任意辅助资料")
        handler = partial(QuietHandler, directory=directory)
        server = ThreadingHTTPServer(("0.0.0.0", FIXTURE_PORT), handler)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        try:
            with httpx.Client(base_url=BASE_URL, timeout=10, trust_env=False) as client:
                response = client.post("/api/v1/draft-reviews", json=payload())
                response.raise_for_status()
                assert response.status_code == 202
                task_id = response.json()["data"]["task_id"]
                history: list[tuple[str, str, int]] = []
                deadline = time.monotonic() + 40
                while time.monotonic() < deadline:
                    detail = client.get(f"/api/v1/tasks/{task_id}")
                    detail.raise_for_status()
                    data = detail.json()["data"]
                    state = (data["status"], data["stage"], data["progress"])
                    if not history or history[-1] != state:
                        history.append(state)
                    if data["status"] in {"SUCCEEDED", "FAILED"}:
                        break
                    time.sleep(0.2)
                assert data["status"] == "SUCCEEDED", (task_id, data)
                result = client.get(f"/api/v1/tasks/{task_id}/result")
                result.raise_for_status()
                result_data = result.json()["data"]
                assert result_data["schema_version"] == "2.1"
                assert result_data["mock"] is False
                assert result_data["metadata"]["execution_mode"] == "RULE_BASED"
                assert result_data["metadata"]["workflow_version"] == "0.5.1"
                assert result_data["metadata"]["rules_version"] == "0.4.1"
                assert result_data["conclusion"] == "PASS"
                assert result_data["summary"]["statistics"]["risk_count"] == 0
                assert len(result_data["files"]) == 3
                assert all(
                    item["parser_name"] == "python-docx" for item in result_data["files"]
                )
                assert (
                    result_data["warnings"][-1]["code"]
                    == "DRAFT_REVIEW_RULE_BASED_LIMITATION"
                )
                print({"task_id": task_id, "history": history, "result": "valid-template-check"})
        finally:
            server.shutdown()
            server.server_close()
            thread.join(timeout=5)


if __name__ == "__main__":
    main()
