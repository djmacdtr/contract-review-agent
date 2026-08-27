from __future__ import annotations

import re

import pdfplumber
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.errors import WorkflowError
from app.documents.models import (
    DocumentBlock,
    DocumentLocation,
    ParsedDocument,
    ParsedTable,
    ProcessingWarning,
    TableCell,
    TableRow,
)
from app.documents.normalization import normalize_text
from app.services.downloader import DOCX_MIME, PDF_MIME, LocalFile

CLAUSE_PATTERN = re.compile(r"^(第[一二三四五六七八九十百千万0-9]+条|\d+(?:\.\d+)*[、.])")


class DocxParser:
    name = "python-docx"

    async def parse(self, file: LocalFile) -> ParsedDocument:
        try:
            document = Document(file.path)
            blocks: list[DocumentBlock] = []
            warnings: list[ProcessingWarning] = []
            paragraph_index = 0
            table_index = 0
            order = 0
            for child in document.element.body.iterchildren():
                if isinstance(child, CT_P):
                    paragraph = Paragraph(child, document)
                    raw = paragraph.text.strip()
                    if not raw:
                        continue
                    section = (
                        raw
                        if paragraph.style.name.startswith("Heading")
                        or CLAUSE_PATTERN.match(raw)
                        else None
                    )
                    blocks.append(
                        DocumentBlock(
                            block_id=f"{file.file_id}_p{paragraph_index:06d}",
                            type="PARAGRAPH",
                            order=order,
                            raw_text=raw,
                            normalized_text=normalize_text(raw),
                            location=DocumentLocation(
                                paragraph_index=paragraph_index,
                                section=section,
                                structure_id=f"paragraph:{paragraph_index}",
                            ),
                        )
                    )
                    paragraph_index += 1
                    order += 1
                elif isinstance(child, CT_Tbl):
                    table = Table(child, document)
                    rows: list[TableRow] = []
                    merged = False
                    for row_index, row in enumerate(table.rows):
                        seen_cells: set[int] = set()
                        cells: list[TableCell] = []
                        for column, cell in enumerate(row.cells):
                            identity = id(cell._tc)
                            merged = merged or identity in seen_cells
                            seen_cells.add(identity)
                            raw = normalize_text(cell.text)
                            cells.append(
                                TableCell(
                                    raw_text=raw,
                                    normalized_text=normalize_text(raw),
                                    location=DocumentLocation(
                                        table_index=table_index,
                                        row=row_index,
                                        column=column,
                                        structure_id=(
                                            f"table_cell:{table_index}:{row_index}:{column}"
                                        ),
                                    ),
                                )
                            )
                        rows.append(TableRow(row=row_index, cells=cells))
                    parsed_table = ParsedTable(table_index=table_index, rows=rows)
                    raw_table = "\n".join(
                        "\t".join(cell.raw_text for cell in row.cells) for row in rows
                    )
                    blocks.append(
                        DocumentBlock(
                            block_id=f"{file.file_id}_t{table_index:06d}",
                            type="TABLE",
                            order=order,
                            raw_text=raw_table,
                            normalized_text=normalize_text(raw_table),
                            location=DocumentLocation(
                                table_index=table_index,
                                structure_id=f"table:{table_index}",
                            ),
                            table=parsed_table,
                        )
                    )
                    if merged:
                        warnings.append(
                            ProcessingWarning(
                                code="DOCX_MERGED_CELLS_SIMPLIFIED",
                                message=(
                                    f"表格 {table_index} 包含合并单元格，"
                                    "已按可见单元格简化处理"
                                ),
                            )
                        )
                    table_index += 1
                    order += 1
            return ParsedDocument(
                file_id=file.file_id,
                role=file.role,
                file_name=file.file_name,
                sha256=file.sha256,
                page_count=None,
                blocks=blocks,
                parser_name=self.name,
                parser_metadata={"physical_page_numbers": False},
                warnings=warnings,
            )
        except WorkflowError:
            raise
        except Exception as exc:
            raise WorkflowError("PARSE_FAILED", "DOCX 文档解析失败") from exc


class TextPdfParser:
    name = "pdfplumber"

    def __init__(self, *, min_text_chars_per_page: int = 20) -> None:
        self.min_text_chars_per_page = min_text_chars_per_page

    async def parse(self, file: LocalFile) -> ParsedDocument:
        try:
            blocks: list[DocumentBlock] = []
            total_chars = 0
            order = 0
            with pdfplumber.open(file.path) as pdf:
                page_count = len(pdf.pages)
                for page_number, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    total_chars += len(re.sub(r"\s+", "", text))
                    for line_index, raw_line in enumerate(text.splitlines()):
                        raw = raw_line.strip()
                        if not raw:
                            continue
                        blocks.append(
                            DocumentBlock(
                                block_id=f"{file.file_id}_page{page_number:04d}_line{line_index:05d}",
                                type="PARAGRAPH",
                                order=order,
                                raw_text=raw,
                                normalized_text=normalize_text(raw),
                                location=DocumentLocation(
                                    page=page_number, paragraph_index=line_index
                                ),
                            )
                        )
                        order += 1
            if total_chars < max(1, page_count) * self.min_text_chars_per_page:
                raise WorkflowError("OCR_REQUIRED", "PDF 文本层为空或过少，需要 OCR 处理")
            return ParsedDocument(
                file_id=file.file_id,
                role=file.role,
                file_name=file.file_name,
                sha256=file.sha256,
                page_count=page_count,
                blocks=blocks,
                parser_name=self.name,
                parser_metadata={
                    "pdf_tables_structured": False,
                    "physical_page_numbers": True,
                    "text_chars": total_chars,
                },
            )
        except WorkflowError:
            raise
        except Exception as exc:
            raise WorkflowError("PARSE_FAILED", "PDF 文档解析失败") from exc


class ParserRegistry:
    def __init__(self, *, pdf_min_text_chars_per_page: int = 20) -> None:
        self.docx = DocxParser()
        self.pdf = TextPdfParser(min_text_chars_per_page=pdf_min_text_chars_per_page)

    async def parse(self, file: LocalFile) -> ParsedDocument:
        if file.detected_mime_type == DOCX_MIME:
            return await self.docx.parse(file)
        if file.detected_mime_type == PDF_MIME:
            return await self.pdf.parse(file)
        raise WorkflowError("UNSUPPORTED_FILE_TYPE", "没有可用的文档解析器")
