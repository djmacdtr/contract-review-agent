from pathlib import Path

import httpx
from docx import Document
from reportlab.pdfgen.canvas import Canvas

from app.comparison.engine import CompareOptions, compare_documents
from app.core.config import Settings
from app.core.enums import TaskStage, TaskType
from app.documents.models import DocumentBlock, DocumentLocation, ParsedDocument, ProcessingWarning
from app.documents.parsers import ParserRegistry
from app.documents.router import DocumentParsingRouter
from app.services.downloader import SafeFileDownloadService
from app.workflows.final_compare import FinalCompareWorkflowExecutor


def build_docx(path: Path, amount: str) -> bytes:
    document = Document()
    document.add_heading("第一条 合同金额", level=1)
    document.add_paragraph(f"本合同金额为{amount}万元，期限为24个月。")
    document.save(path)
    return path.read_bytes()


async def resolver(host: str, port: int) -> list[str]:
    return ["127.0.0.1"]


async def test_final_compare_graph_returns_rule_based_traceable_result_and_cleans(
    tmp_path: Path,
) -> None:
    baseline_bytes = build_docx(tmp_path / "baseline.docx", "100")
    target_bytes = build_docx(tmp_path / "target.docx", "120")

    async def handler(request: httpx.Request) -> httpx.Response:
        body = baseline_bytes if request.url.path.endswith("baseline.docx") else target_bytes
        return httpx.Response(200, content=body, request=request)

    settings = Settings(
        _env_file=None,
        TEMP_ROOT=str(tmp_path / "workspaces"),
        ALLOW_HTTP_DOWNLOADS=True,
        DOWNLOAD_HOST_ALLOWLIST="fixture-server",
        MOCK_STAGE_DELAY_SECONDS=0,
    )
    downloader = SafeFileDownloadService(
        settings,
        transport=httpx.MockTransport(handler),
        resolver=resolver,
    )
    executor = FinalCompareWorkflowExecutor(settings, downloader=downloader)
    updates: list[tuple[TaskStage, int]] = []

    async def progress(stage: TaskStage, value: int, message: str) -> None:
        updates.append((stage, value))

    output = await executor.run(
        task_id="tsk_real_compare",
        task_type=TaskType.FINAL_COMPARE,
        files=[
            {
                "file_id": "fil_base",
                "role": "BASELINE",
                "file_name": "baseline.docx",
                "url": "http://fixture-server/baseline.docx?token=secret",
                "safe_url": "http://fixture-server/baseline.docx",
            },
            {
                "file_id": "fil_target",
                "role": "TARGET",
                "file_name": "target.docx",
                "url": "http://fixture-server/target.docx?token=secret",
                "safe_url": "http://fixture-server/target.docx",
            },
        ],
        options={"numeric_sensitive": True, "ignore_headers_footers": True},
        progress_callback=progress,
    )
    result = output.result
    assert result["mock"] is False
    assert result["metadata"]["execution_mode"] == "RULE_BASED"
    assert result["metadata"]["workflow_version"] == "0.4.1"
    assert result["metadata"]["rules_version"] == "0.4.1"
    assert result["metadata"]["comparison_diagnostics"]["reliable"] is True
    assert result["metadata"]["primary_model"] is None
    assert result["metadata"]["model_runs"] == []
    assert result["conclusion"] == "RISK_FOUND"
    assert any(item["diff_type"] == "NUMERIC_CHANGED" for item in result["diff_items"])
    assert result["files"][0]["sha256"] == output.file_metadata[0]["sha256"]
    assert "token=secret" not in str(result)
    assert updates[-1][0] == TaskStage.PERSISTING_RESULT
    assert not any((tmp_path / "workspaces").iterdir())


class SyntheticOcrParser:
    async def parse(self, file, *, mode: str) -> ParsedDocument:
        amount = "100" if file.role == "BASELINE" else "120"
        return ParsedDocument(
            file_id=file.file_id,
            role=file.role,
            file_name=file.file_name,
            sha256=file.sha256,
            page_count=1,
            blocks=[
                DocumentBlock(
                    block_id=f"{file.file_id}_ocr_1",
                    type="PARAGRAPH",
                    order=0,
                    raw_text=f"合同金额为{amount}万元。",
                    normalized_text=f"合同金额为{amount}万元。",
                    location=DocumentLocation(page=1, source="OCR", confidence=0.99),
                )
            ],
            parser_name="textin-document-parser",
            parser_metadata={"ocr": True, "engine_version": "test-engine", "parse_mode": mode},
            warnings=[
                ProcessingWarning(
                    code="OCR_USED",
                    message="文档由外部 OCR 服务解析",
                    requires_manual_review=False,
                )
            ],
        )


async def test_scan_pdf_graph_uses_external_parser_and_exposes_metadata(tmp_path: Path) -> None:
    empty_pdf_path = tmp_path / "scan.pdf"
    canvas = Canvas(str(empty_pdf_path))
    canvas.showPage()
    canvas.save()
    empty_pdf = empty_pdf_path.read_bytes()

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=empty_pdf, request=request)

    settings = Settings(
        _env_file=None,
        TEMP_ROOT=str(tmp_path / "workspaces"),
        ALLOW_HTTP_DOWNLOADS=True,
        DOWNLOAD_HOST_ALLOWLIST="fixture-server",
        OCR_ENABLED=True,
        MOCK_STAGE_DELAY_SECONDS=0,
    )
    downloader = SafeFileDownloadService(
        settings,
        transport=httpx.MockTransport(handler),
        resolver=resolver,
    )
    document_router = DocumentParsingRouter(
        local=ParserRegistry(pdf_min_text_chars_per_page=20),
        external=SyntheticOcrParser(),
    )
    executor = FinalCompareWorkflowExecutor(
        settings,
        downloader=downloader,
        document_router=document_router,
    )

    async def progress(stage: TaskStage, value: int, message: str) -> None:
        return None

    output = await executor.run(
        task_id="tsk_ocr_compare",
        task_type=TaskType.FINAL_COMPARE,
        files=[
            {
                "file_id": "fil_base",
                "role": "BASELINE",
                "file_name": "base.pdf",
                "url": "http://fixture-server/base.pdf",
                "safe_url": "http://fixture-server/base.pdf",
            },
            {
                "file_id": "fil_target",
                "role": "TARGET",
                "file_name": "target.pdf",
                "url": "http://fixture-server/target.pdf",
                "safe_url": "http://fixture-server/target.pdf",
            },
        ],
        options={},
        progress_callback=progress,
    )
    assert output.result["conclusion"] == "RISK_FOUND"
    assert output.result["files"][0]["parser_metadata"]["engine_version"] == "test-engine"
    assert output.result["diff_items"][0]["baseline"]["location"]["source"] == "OCR"
    assert "包含 OCR 解析" in output.result["warnings"][-1]["message"]
    assert not any((tmp_path / "workspaces").iterdir())


def test_only_low_confidence_ocr_text_diffs_require_review_instead_of_risk() -> None:
    settings = Settings(_env_file=None, OCR_LOW_CONFIDENCE_THRESHOLD=0.8)
    documents = []
    for file_id, role, text in (
        ("fil_base", "BASELINE", "第一条 普通说明。"),
        ("fil_target", "TARGET", "第一条 普通描述。"),
    ):
        documents.append(
            ParsedDocument(
                file_id=file_id,
                role=role,
                file_name=f"{role.lower()}.pdf",
                sha256="d" * 64,
                page_count=1,
                blocks=[
                    DocumentBlock(
                        block_id=f"{file_id}_p1",
                        type="PARAGRAPH",
                        order=0,
                        raw_text=text,
                        normalized_text=text,
                        location=DocumentLocation(page=1, source="OCR", confidence=0.55),
                    )
                ],
                parser_name="textin-document-parser",
                parser_metadata={"ocr": True},
                warnings=[
                    ProcessingWarning(
                        code="OCR_LOW_CONFIDENCE",
                        message="OCR 置信度低，需要人工复核",
                        confidence=0.55,
                    )
                ],
            )
        )
    comparison = compare_documents(
        documents[0],
        documents[1],
        CompareOptions(ocr_low_confidence_threshold=0.8),
    )
    executor = FinalCompareWorkflowExecutor(settings)
    result = executor._build_result(
        "tsk_low_confidence",
        [
            {"file_id": "fil_base", "safe_url": "https://example.test/base.pdf"},
            {"file_id": "fil_target", "safe_url": "https://example.test/target.pdf"},
        ],
        documents,
        comparison,
    )
    assert result["conclusion"] == "REVIEW_REQUIRED"
    assert result["diff_items"][0]["severity"] == "LOW"
    assert result["diff_items"][0]["review_reason"] == "OCR_LOW_CONFIDENCE_VARIANCE"
