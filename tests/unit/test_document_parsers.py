import asyncio
import time
from pathlib import Path

import pytest
from docx import Document
from reportlab.pdfgen.canvas import Canvas

from app.core.errors import WorkflowError
from app.documents.parsers import DocxParser, TextPdfParser
from app.services.downloader import LocalFile


def local_file(path: Path, *, mime: str, file_id: str = "fil_1") -> LocalFile:
    return LocalFile(
        file_id=file_id,
        role="BASELINE",
        file_name=path.name,
        safe_url="https://files.example.com/safe",
        path=path,
        file_size=path.stat().st_size,
        sha256="a" * 64,
        detected_mime_type=mime,
    )


async def test_docx_parser_preserves_paragraph_table_order_and_locations(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("第一条 合同金额", level=1)
    document.add_paragraph("融资金额为100万元。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "金额"
    table.cell(1, 0).text = "设备A"
    table.cell(1, 1).text = "100万元"
    document.add_paragraph("合同期限为24个月。")
    document.save(path)

    parsed = await DocxParser().parse(
        local_file(
            path,
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
    )
    assert [block.type for block in parsed.blocks] == [
        "PARAGRAPH",
        "PARAGRAPH",
        "TABLE",
        "PARAGRAPH",
    ]
    table_block = parsed.blocks[2]
    assert table_block.table is not None
    assert table_block.table.rows[1].cells[1].raw_text == "100万元"
    assert table_block.table.rows[1].cells[1].location.table_index == 0
    assert parsed.page_count is None


async def test_text_pdf_parser_has_page_locations_and_rejects_empty_pdf(tmp_path: Path) -> None:
    text_path = tmp_path / "text.pdf"
    canvas = Canvas(str(text_path))
    canvas.drawString(72, 760, "Contract amount is CNY 1000000 and term is 24 months.")
    canvas.showPage()
    canvas.drawString(72, 760, "Second page has deterministic searchable text.")
    canvas.save()
    parsed = await TextPdfParser(min_text_chars_per_page=10).parse(
        local_file(text_path, mime="application/pdf")
    )
    assert parsed.page_count == 2
    assert {block.location.page for block in parsed.blocks} == {1, 2}
    assert parsed.parser_metadata["physical_page_numbers"] is True

    empty_path = tmp_path / "empty.pdf"
    empty = Canvas(str(empty_path))
    empty.showPage()
    empty.save()
    with pytest.raises(WorkflowError) as caught:
        await TextPdfParser(min_text_chars_per_page=10).parse(
            local_file(empty_path, mime="application/pdf", file_id="fil_empty")
        )
    assert caught.value.code == "OCR_REQUIRED"


async def test_local_parser_does_not_block_asyncio_heartbeat(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = tmp_path / "slow.docx"
    document = Document()
    document.add_paragraph("测试内容")
    document.save(path)
    parser = DocxParser()
    original = parser._parse_sync

    def slow_parse(file: LocalFile):
        time.sleep(0.05)
        return original(file)

    monkeypatch.setattr(parser, "_parse_sync", slow_parse)
    ticks = 0
    running = True

    async def heartbeat() -> None:
        nonlocal ticks
        while running:
            ticks += 1
            await asyncio.sleep(0.005)

    task = asyncio.create_task(heartbeat())
    try:
        await parser.parse(
            local_file(
                path,
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
        )
    finally:
        running = False
        await task

    assert ticks >= 5
