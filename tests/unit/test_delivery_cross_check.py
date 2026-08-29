from pathlib import Path
from types import SimpleNamespace

import httpx
import pytest
from docx import Document

from app.adapters.llm.base import LlmResult
from app.core.config import Settings
from app.core.enums import TaskType
from app.core.errors import WorkflowError
from app.documents.models import DocumentBlock, DocumentLocation, ParsedDocument
from app.draft_review.delivery_cross_check import (
    build_reference_candidate_groups,
    candidate_to_fact_matrix_candidate,
    extract_document_candidates,
    group_decision_to_results,
    normalize_candidate_value,
)
from app.draft_review.template_checks import analyze_template
from app.schemas.results import TaskResultData
from app.services.downloader import SafeFileDownloadService
from app.workflows.draft_review import DraftReviewWorkflowExecutor


def document(file_id: str, role: str, text: str) -> ParsedDocument:
    return ParsedDocument(
        file_id=file_id,
        role=role,
        file_name=f"{role.lower()}.docx",
        sha256=(file_id + "0" * 64)[:64],
        page_count=None,
        blocks=[
            DocumentBlock(
                block_id=f"{file_id}_block",
                type="PARAGRAPH",
                order=0,
                raw_text=text,
                normalized_text=text,
                location=DocumentLocation(paragraph_index=0),
            )
        ],
        parser_name="python-docx",
    )


def empty_review():
    return SimpleNamespace(diff_items=[], diagnostics=SimpleNamespace(filtered_diff_items=[]))


def write_docx(path: Path, title: str, body: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    doc.save(path)
    return path.read_bytes()


def test_candidate_normalization_keeps_type_and_removes_presentation_noise() -> None:
    assert normalize_candidate_value("1,000 万元", "MONEY") == "1000万元"
    assert normalize_candidate_value("5％", "PERCENTAGE") == "5%"
    assert normalize_candidate_value("2026-08-27", "DATE") == "2026-08-27"


def test_candidates_preserve_context_and_extract_open_world_value_types() -> None:
    items = extract_document_candidates(
        document(
            "target",
            "TARGET",
            "项目编号：PRJ-2026-001，融资金额1,000万元，比例5%，期限12个月，签署日期2026年8月27日。",
        )
    )
    kinds = {item["value_type"] for item in items}
    assert {"IDENTIFIER", "MONEY", "PERCENTAGE", "DURATION", "DATE"} <= kinds
    assert all(item["context"] for item in items)
    assert all(item["location"]["paragraph_index"] == 0 for item in items)


def test_reference_candidates_are_deduplicated_and_limited_to_three_per_file() -> None:
    target = document("target", "TARGET", "融资金额100万元，期限12个月。")
    reference = document(
        "reference",
        "REFERENCE",
        "融资金额100万元，融资金额100万元，融资金额120万元，融资金额130万元。",
    )
    groups = build_reference_candidate_groups(target, [reference], empty_review())
    assert groups["group_count"] >= 2
    assert all(
        len(values) <= 3
        for group in groups["groups"]
        for values in group["references"].values()
    )


def test_four_decisions_only_materialize_safe_results() -> None:
    target = document("target", "TARGET", "融资金额100万元。")
    reference = document("reference", "REFERENCE", "融资金额120万元。")
    groups = build_reference_candidate_groups(target, [reference], empty_review())["groups"]
    assert groups
    decisions = {
        groups[0]["candidate_id"]: {"decision": "CONFLICT"},
    }
    result = group_decision_to_results(groups, decisions)
    assert len(result["diff_items"]) == 1
    assert result["passed_checks"] == []
    assert result["fact_matrix"] == []

    target_candidate = groups[0]["target"]
    reference_candidate = next(iter(groups[0]["references"].values()))[0]
    matrix_candidate = candidate_to_fact_matrix_candidate(target_candidate)
    assert matrix_candidate["source_file_id"] == "target"
    assert reference_candidate["file_id"] == "reference"


def test_match_creates_fact_matrix_and_pass_only_for_equal_value() -> None:
    target = document("target", "TARGET", "融资金额100万元。")
    reference = document("reference", "REFERENCE", "融资金额100万元。")
    groups = build_reference_candidate_groups(target, [reference], empty_review())["groups"]
    result = group_decision_to_results(
        groups,
        {group["candidate_id"]: {"decision": "MATCH"} for group in groups},
    )
    assert result["fact_matrix"]
    assert result["passed_checks"]
    assert result["diff_items"] == []


def test_template_delta_can_be_used_as_a_candidate_source() -> None:
    template = document("template", "TEMPLATE", "第一条 融资金额为##{融资金额}。")
    target = document("target", "TARGET", "第一条 融资金额为100万元。")
    review = analyze_template(template, target)
    reference = document("reference", "REFERENCE", "第一条 融资金额为120万元。")
    groups = build_reference_candidate_groups(target, [reference], review)
    assert any(group["diff_ids"] for group in groups["groups"])


def test_kiss_result_keeps_template_risks_when_cross_batch_is_unavailable() -> None:
    template = document("template", "TEMPLATE", "第一条 固定条款。")
    target = document("target", "TARGET", "第一条 变更条款。")
    reference = document("reference", "REFERENCE", "辅助资料无数值。")
    review = analyze_template(template, target)
    executor = DraftReviewWorkflowExecutor(Settings(_env_file=None))
    result = executor._build_delivery_result(
        "task",
        [
            {"file_id": "target", "safe_url": "http://fixture/target.docx"},
            {"file_id": "template", "safe_url": "http://fixture/template.docx"},
            {"file_id": "reference", "safe_url": "http://fixture/reference.docx"},
        ],
        [target, template, reference],
        review,
        {"groups": [], "warnings": [], "group_count": 0},
        {
            "decisions": {},
            "warnings": [
                {
                    "code": "LLM_CROSS_VALIDATE_UNAVAILABLE",
                    "message": "跨资料对应判断未完成。",
                    "requires_manual_review": False,
                }
            ],
            "model_runs": [],
        },
        {},
    )
    validated = TaskResultData.model_validate(result)
    assert validated.review_items == []
    assert result["risk_items"]
    assert all(item["analysis_advice"] for item in result["risk_items"])
    assert result["fact_matrix"] == []


def test_default_graph_uses_full_fact_extraction_chain() -> None:
    executor = DraftReviewWorkflowExecutor(Settings(_env_file=None))

    async def progress(*_args):
        return None

    graph = executor._build_graph(None, progress)  # type: ignore[arg-type]
    graph_nodes = set(graph.get_graph().nodes)
    assert "build_reference_candidates" not in graph_nodes
    assert "cross_validate_candidates" not in graph_nodes
    assert "extract_facts" in graph_nodes
    assert "map_cross_document_facts" in graph_nodes
    assert "plan_semantics" in graph_nodes


class DeliveryLlm:
    def __init__(self) -> None:
        self.cross_calls = 0
        self.advice_calls = 0

    async def cross_validate_candidates(self, payload: dict) -> LlmResult:
        self.cross_calls += 1
        items = []
        for group in payload["candidates"]:
            target = group["target"]
            references = [
                item
                for values in group["references"].values()
                for item in values
            ]
            items.append(
                {
                    "candidate_id": group["candidate_id"],
                    "decision": (
                        "MATCH"
                        if any(
                            item["normalized_value"] == target["normalized_value"]
                            for item in references
                        )
                        else "CONFLICT"
                    ),
                    "reason": "测试判断",
                }
            )
        return LlmResult(
            value={"items": items},
            configured_model="delivery-test",
            actual_model="delivery-test",
            mock=False,
        )

    async def generate_delivery_advice(self, payload: dict) -> LlmResult:
        self.advice_calls += 1
        return LlmResult(
            value={
                "overall_advice": "请依据证据核对风险。",
                "priority_actions": [],
                "manual_review_focus": [],
                "limitations": [],
                "risk_advices": [
                    {
                        "risk_id": risk["risk_id"],
                        "analysis_advice": "请根据对应文件证据核对该项差异。",
                    }
                    for risk in payload.get("risk_items", [])
                ],
            },
            configured_model="delivery-test",
            actual_model="delivery-test",
            mock=False,
        )


@pytest.mark.asyncio
async def test_kiss_only_llm_cannot_soft_degrade_the_delivery_graph(
    tmp_path: Path,
) -> None:
    bodies = {
        "/target.docx": write_docx(tmp_path / "target.docx", "合同", "融资金额为100万元。"),
        "/template.docx": write_docx(
            tmp_path / "template.docx", "合同", "融资金额为##{融资金额}万元。"
        ),
        "/reference.docx": write_docx(
            tmp_path / "reference.docx", "项目资料", "融资金额为120万元。"
        ),
    }

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=bodies[request.url.path], request=request)

    async def resolver(_host: str, _port: int) -> list[str]:
        return ["127.0.0.1"]

    settings = Settings(
        _env_file=None,
        TEMP_ROOT=str(tmp_path / "workspaces"),
        ALLOW_HTTP_DOWNLOADS=True,
        DOWNLOAD_HOST_ALLOWLIST="fixture-server",
        LLM_ENABLED=True,
        LLM_BASE_URL="https://llm.invalid",
        LLM_API_KEY="not-used",
    )
    llm = DeliveryLlm()
    executor = DraftReviewWorkflowExecutor(
        settings,
        downloader=SafeFileDownloadService(
            settings,
            transport=httpx.MockTransport(handler),
            resolver=resolver,
        ),
        llm=llm,
    )
    with pytest.raises(WorkflowError) as caught:
        await executor.run(
            task_id="task_delivery",
            task_type=TaskType.DRAFT_REVIEW,
            files=[
            {
                "file_id": "target",
                "role": "TARGET",
                "file_name": "target.docx",
                "url": "http://fixture-server/target.docx",
                "safe_url": "http://fixture-server/target.docx",
            },
            {
                "file_id": "template",
                "role": "TEMPLATE",
                "file_name": "template.docx",
                "url": "http://fixture-server/template.docx",
                "safe_url": "http://fixture-server/template.docx",
            },
            {
                "file_id": "reference",
                "role": "REFERENCE",
                "file_name": "reference.docx",
                "url": "http://fixture-server/reference.docx",
                "safe_url": "http://fixture-server/reference.docx",
            },
            ],
            options={},
            progress_callback=lambda *_args: _noop_progress(),
        )
    assert caught.value.details["failure_code"] == "FACT_EXTRACTION_METHOD_UNAVAILABLE"
    assert llm.cross_calls == 0
    assert llm.advice_calls == 0


async def _noop_progress() -> None:
    return None
