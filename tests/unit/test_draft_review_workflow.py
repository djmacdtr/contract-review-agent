from pathlib import Path

import httpx
import pytest
from docx import Document

from app.adapters.llm.base import LlmResult
from app.adapters.llm.openai_client import LlmClientError
from app.core.config import Settings
from app.core.enums import TaskStage, TaskType
from app.core.errors import WorkflowError
from app.schemas.results import TaskResultData
from app.services.downloader import SafeFileDownloadService
from app.workflows.draft_review import (
    DraftReviewWorkflowExecutor,
    _dynamic_failure_code,
    logger,
)


class EvidencedLlm:
    async def probe_models(self) -> list[str]:
        return ["test-model"]

    async def generate_advice(self, payload: dict) -> LlmResult:
        raise AssertionError("advice is not used in this workflow slice")

    async def extract_facts(self, payload: dict) -> LlmResult:
        first = payload["blocks"][0]
        file_id = payload["file_id"]
        return LlmResult(
            value={
                "profile": {
                    "file_id": file_id,
                    "document_kind": "合同" if payload["role"] == "TARGET" else "辅助资料",
                    "title": None,
                    "confidence": 0.8,
                    "evidence_locations": [first["location"]],
                },
                "facts": [
                    {
                        "field_key": "document_title",
                        "display_name": "文档标题",
                        "value_type": "TEXT",
                        "raw_value": first["text"],
                        "normalized_hint": None,
                        "source_file_id": file_id,
                        "evidence_text": first["text"],
                        "location": first["location"],
                        "confidence": 0.8,
                    }
                ],
                "missing_field_keys": [],
            },
            configured_model="test-model",
            actual_model="test-model-v1",
            mock=False,
            duration_ms=10,
            request_attempts=1,
        )


class MapReduceFixtureLlm:
    """Small new-interface fixture that exercises the Send-based extractor."""

    def __init__(self) -> None:
        self.profile_calls = 0
        self.batch_payloads: list[dict] = []

    async def probe_models(self) -> list[str]:
        return ["map-reduce-fixture"]

    async def extract_document_profile(self, payload: dict) -> LlmResult:
        self.profile_calls += 1
        first = payload["overview_blocks"][0]
        return LlmResult(
            value={
                "document_kind": "合成资料",
                "title": None,
                "confidence": 0.9,
                "evidence_locations": [first["location"]],
            },
            configured_model="map-reduce-fixture",
            actual_model="map-reduce-fixture",
            mock=False,
        )

    async def extract_fact_batch(self, payload: dict) -> LlmResult:
        self.batch_payloads.append(payload)
        return LlmResult(
            value={
                "facts": [],
                "numeric_candidate_decisions": [
                    {
                        "candidate_index": candidate["candidate_index"],
                        "decision": "IGNORE",
                        "reason_code": "FIXTURE_IGNORE",
                    }
                    for candidate in payload["numeric_candidates"]
                ],
            },
            configured_model="map-reduce-fixture",
            actual_model="map-reduce-fixture",
            mock=False,
        )


class SplittingMapReduceFixtureLlm(MapReduceFixtureLlm):
    async def extract_fact_batch(self, payload: dict) -> LlmResult:
        if len(payload["units"]) > 1:
            self.batch_payloads.append(payload)
            raise LlmClientError("LLM_INVALID_JSON", "模型未返回有效 JSON")
        return await super().extract_fact_batch(payload)


class ClassifiedFailureLlm(EvidencedLlm):
    async def extract_facts(self, payload: dict) -> LlmResult:
        raise LlmClientError(
            "LLM_EXTRACTION_EVIDENCE_INVALID",
            "模型紧凑事实结果未通过安全证据校验",
            failure_code="FACT_LOCATION_NOT_FOUND",
        )


class SchemaFailureLlm(EvidencedLlm):
    async def extract_facts(self, payload: dict) -> LlmResult:
        raise LlmClientError(
            "LLM_SCHEMA_INVALID",
            "模型事实结果不符合结构约束",
            failure_code="LLM_RESPONSE_SCHEMA_INVALID",
            validation_summary={
                "error_count": 1,
                "items": [
                    {
                        "path": "facts.*.value_type",
                        "error_type": "literal_error",
                        "count": 1,
                    }
                ],
                "truncated": False,
            },
        )


class SchemaFailureWithoutSummaryLlm(EvidencedLlm):
    async def extract_facts(self, payload: dict) -> LlmResult:
        raise LlmClientError(
            "LLM_SCHEMA_INVALID",
            "模型事实结果不符合结构约束",
            failure_code="LLM_RESPONSE_SCHEMA_INVALID",
        )


class RawSchemaFailureLlm(EvidencedLlm):
    async def extract_facts(self, payload: dict) -> LlmResult:
        result = await super().extract_facts(payload)
        result.value["facts"][0]["value_type"] = "NOT_A_VALUE"
        return result


class SplitAfterJsonFailureLlm(EvidencedLlm):
    def __init__(self) -> None:
        self.payloads: list[dict] = []
        self.failed = False

    async def extract_facts(self, payload: dict) -> LlmResult:
        self.payloads.append(payload)
        if len(payload["blocks"]) > 1 and not self.failed:
            self.failed = True
            raise LlmClientError("LLM_INVALID_JSON", "模型未返回有效 JSON")
        return await super().extract_facts(payload)


class AlwaysJsonFailureLlm(EvidencedLlm):
    def __init__(self) -> None:
        self.calls = 0

    async def extract_facts(self, payload: dict) -> LlmResult:
        self.calls += 1
        raise LlmClientError("LLM_INVALID_JSON", "模型未返回有效 JSON")


class EnvelopeFailureLlm(EvidencedLlm):
    def __init__(self) -> None:
        self.calls = 0

    async def extract_facts(self, payload: dict) -> LlmResult:
        self.calls += 1
        raise LlmClientError("LLM_RESPONSE_INVALID", "模型响应结构无效")


def build_docx(path: Path, title: str, body: str) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    document.save(path)
    return path.read_bytes()


def build_table_docx(path: Path, title: str, body: str, *, expanded: bool) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "填写值"
    table.cell(1, 0).text = "融资金额"
    table.cell(1, 1).text = "100"
    if expanded:
        row = table.add_row()
        row.cells[0].text = "附加字段"
        row.cells[1].text = "附加内容"
    document.save(path)
    return path.read_bytes()


class ConsensusFixtureLlm:
    def __init__(
        self,
        *,
        extraction_confidence: float = 0.95,
        review_confidence: float = 0.95,
        evidence_complete: bool = True,
        extraction_model: str = "extractor-v1",
        review_model: str = "reviewer-v1",
        review_failure: bool = False,
        include_numeric_rule: bool = False,
        review_rule_mismatch: bool = False,
        invalid_advice_evidence: bool = False,
    ) -> None:
        self.extraction_confidence = extraction_confidence
        self.review_confidence = review_confidence
        self.evidence_complete = evidence_complete
        self.extraction_model = extraction_model
        self.review_model = review_model
        self.review_failure = review_failure
        self.include_numeric_rule = include_numeric_rule
        self.review_rule_mismatch = review_rule_mismatch
        self.invalid_advice_evidence = invalid_advice_evidence
        self.advice_calls = 0
        self.review_payloads: list[dict] = []

    async def probe_models(self) -> list[str]:
        return [self.extraction_model, self.review_model]

    async def extract_facts(self, payload: dict) -> LlmResult:
        first = payload["blocks"][0]
        file_id = payload["file_id"]
        spec = (
            {
                "validation_id": "amount_matches_literal",
                "display_name": "金额校验",
                "expression": {
                    "op": "equals",
                    "left": {"op": "fact", "concept_id": "amount"},
                    "right": {"op": "literal", "value": "100"},
                },
                "evidence_locations": [first["location"]],
                "confidence": self.extraction_confidence,
            }
            if self.include_numeric_rule
            else None
        )
        return LlmResult(
            value={
                "profile": {
                    "file_id": file_id,
                    "document_kind": "合同",
                    "title": None,
                    "confidence": self.extraction_confidence,
                    "evidence_locations": [first["location"]],
                },
                "facts": [
                    {
                        "field_key": "amount",
                        "display_name": "金额",
                        "value_type": "MONEY",
                        "raw_value": "100",
                        "normalized_hint": None,
                        "source_file_id": file_id,
                        "evidence_text": first["text"],
                        "location": first["location"],
                        "confidence": self.extraction_confidence,
                    }
                ],
                "missing_field_keys": [],
                "semantic_concepts": [],
                "validation_specs": [spec] if spec else [],
            },
            configured_model=self.extraction_model,
            actual_model=self.extraction_model,
            mock=False,
        )

    async def review_facts(self, payload: dict) -> LlmResult:
        self.review_payloads.append(payload)
        if self.review_failure:
            raise LlmClientError("LLM_UPSTREAM_ERROR", "模型评审不可用")
        fact = payload["facts"][0]
        validation_specs = payload.get("validation_specs", [])
        if self.review_rule_mismatch and validation_specs:
            validation_specs = [
                {
                    **validation_specs[0],
                    "expression": {
                        **validation_specs[0]["expression"],
                        "right": {"op": "literal", "value": "101"},
                    },
                }
            ]
        return LlmResult(
            value={
                "file_id": payload["file_id"],
                "decisions": [
                    {
                        "field_key": fact["field_key"],
                        "source_file_id": fact["source_file_id"],
                        "location": fact["location"],
                        "decision": "ACCEPT",
                        "evidence_text": fact["evidence_text"],
                        "confidence": self.review_confidence,
                        "reason_code": "EVIDENCE_MATCHED",
                    }
                ],
                "semantic_concepts": [],
                "validation_specs": validation_specs,
                "confidence": self.review_confidence,
                "evidence_complete": self.evidence_complete,
            },
            configured_model=self.review_model,
            actual_model=self.review_model,
            mock=False,
        )

    async def map_facts(self, payload: dict) -> LlmResult:
        fact = payload["reference_facts"][0]
        return LlmResult(
            value={
                "reference_file_id": payload["reference_file_id"],
                "mappings": [
                    {
                        "target_fact_id": payload["target_facts"][0]["target_fact_id"],
                        "reference_field_key": fact["field_key"],
                        "source_file_id": fact["source_file_id"],
                        "reference_location": fact["location"],
                        "decision": "MATCH",
                        "confidence": self.extraction_confidence,
                        "reason_code": "SAME_BUSINESS_FACT",
                    }
                ],
                "missing_requirements": [],
            },
            configured_model=self.extraction_model,
            actual_model=self.extraction_model,
            mock=False,
        )

    async def review_mappings(self, payload: dict) -> LlmResult:
        proposal = payload["proposed_mapping"]["mappings"][0]
        return LlmResult(
            value={
                "reference_file_id": payload["reference_file_id"],
                "decisions": [
                    {
                        "target_fact_id": proposal["target_fact_id"],
                        "reference_field_key": proposal["reference_field_key"],
                        "source_file_id": proposal["source_file_id"],
                        "reference_location": proposal["reference_location"],
                        "decision": "ACCEPT",
                        "confidence": self.review_confidence,
                        "reason_code": "MAPPING_VERIFIED",
                    }
                ],
                "missing_requirement_decisions": [],
                "confidence": self.review_confidence,
                "evidence_complete": self.evidence_complete,
            },
            configured_model=self.review_model,
            actual_model=self.review_model,
            mock=False,
        )

    async def generate_advice(self, payload: dict) -> LlmResult:
        self.advice_calls += 1
        risk_advices = [
            {
                "risk_id": risk["risk_id"],
                "analysis_advice": f"请结合当前文件位置核对{risk['title']}的业务依据。",
            }
            for risk in payload.get("risk_items", [])
        ]
        if self.invalid_advice_evidence:
            risk_advices.append(
                {
                    "risk_id": "risk_unknown",
                    "analysis_advice": "这条建议不属于当前任务。",
                }
            )
        return LlmResult(
            value={
                "overall_advice": "请复核已有证据。",
                "priority_actions": [],
                "manual_review_focus": [],
                "limitations": [],
                "evidence_refs": [],
                "risk_advices": risk_advices,
            },
            configured_model=self.extraction_model,
            actual_model=self.extraction_model,
            mock=False,
        )


class SemanticPlanFixtureLlm(ConsensusFixtureLlm):
    def __init__(self) -> None:
        super().__init__()
        self.plan_calls = 0

    async def plan_semantics(self, payload: dict) -> LlmResult:
        self.plan_calls += 1
        target = next(item for item in payload["documents"] if item["role"] == "TARGET")
        fact = target["facts"][0]
        return LlmResult(
            value={
                "file_id": payload["file_id"],
                "semantic_concepts": [
                    {
                        "concept_id": "amount",
                        "display_name": "金额",
                        "value_type": "MONEY",
                        "aliases": [],
                        "fact_refs": [
                            {
                                "fact_id": fact["fact_id"],
                                "source_file_id": fact["source_file_id"],
                            }
                        ],
                        "evidence_refs": [
                            {
                                "source_file_id": fact["source_file_id"],
                                "location": fact["location"],
                            }
                        ],
                        "confidence": 0.95,
                    }
                ],
                "validation_specs": [
                    {
                        "validation_id": "amount_positive",
                        "display_name": "金额为正",
                        "expression": {
                            "op": "greater_than",
                            "left": {
                                "op": "fact",
                                "fact_id": fact["fact_id"],
                                "source_file_id": fact["source_file_id"],
                            },
                            "right": {"op": "literal", "value": "0"},
                        },
                        "evidence_refs": [
                            {
                                "source_file_id": fact["source_file_id"],
                                "location": fact["location"],
                            }
                        ],
                        "confidence": 0.95,
                    }
                ],
            },
            configured_model=self.extraction_model,
            actual_model=self.extraction_model,
            mock=False,
        )


async def run_consensus_fixture(
    tmp_path: Path,
    llm: ConsensusFixtureLlm,
    *,
    options: dict | None = None,
    same_model_diagnostic: bool = False,
    target_body: str = "金额100",
    template_body: str = "金额100",
    expanded_table: bool = False,
    settings_overrides: dict | None = None,
) -> dict:
    target_bytes = (
        build_table_docx(
            tmp_path / "target.docx", "合同", target_body, expanded=True
        )
        if expanded_table
        else build_docx(tmp_path / "target.docx", "合同", target_body)
    )
    template_bytes = (
        build_table_docx(
            tmp_path / "template.docx", "合同", template_body, expanded=False
        )
        if expanded_table
        else build_docx(tmp_path / "template.docx", "合同", template_body)
    )
    bodies = {
        "/target.docx": target_bytes,
        "/template.docx": template_bytes,
        "/reference.docx": build_docx(tmp_path / "reference.docx", "资料", "金额100"),
    }

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=bodies[request.url.path], request=request)

    async def no_progress(*_args) -> None:
        return None

    settings = Settings(
        _env_file=None,
        TEMP_ROOT=str(tmp_path / "workspaces"),
        ALLOW_HTTP_DOWNLOADS=True,
        DOWNLOAD_HOST_ALLOWLIST="fixture-server",
        LLM_ENABLED=True,
        LLM_BASE_URL="https://llm.invalid",
        LLM_API_KEY="unused",
        LLM_SAME_MODEL_DIAGNOSTIC=same_model_diagnostic,
        **(settings_overrides or {}),
    )
    executor = DraftReviewWorkflowExecutor(
        settings,
        downloader=SafeFileDownloadService(
            settings, transport=httpx.MockTransport(handler), resolver=resolver
        ),
        llm=llm,
    )
    output = await executor.run(
        task_id="tsk_consensus_fixture",
        task_type=TaskType.DRAFT_REVIEW,
        files=[
            {
                "file_id": "fil_target",
                "role": "TARGET",
                "file_name": "target.docx",
                "url": "http://fixture-server/target.docx",
                "safe_url": "http://fixture-server/target.docx",
            },
            {
                "file_id": "fil_template",
                "role": "TEMPLATE",
                "file_name": "template.docx",
                "url": "http://fixture-server/template.docx",
                "safe_url": "http://fixture-server/template.docx",
            },
            {
                "file_id": "fil_reference",
                "role": "REFERENCE",
                "file_name": "reference.docx",
                "url": "http://fixture-server/reference.docx",
                "safe_url": "http://fixture-server/reference.docx",
            },
        ],
        options=options or {},
        progress_callback=no_progress,
    )
    return output.result


async def resolver(host: str, port: int) -> list[str]:
    return ["127.0.0.1"]


async def test_new_send_map_reduce_path_extracts_profile_once_and_reduces_batches(
    tmp_path: Path,
) -> None:
    llm = MapReduceFixtureLlm()

    result = await run_consensus_fixture(tmp_path, llm)  # type: ignore[arg-type]

    TaskResultData.model_validate(result)
    assert llm.profile_calls == 3
    assert llm.batch_payloads
    assert all("profile" not in payload for payload in llm.batch_payloads)
    assert all(
        "source_file_id" not in unit
        for payload in llm.batch_payloads
        for unit in payload["units"]
    )
    assert result["metadata"]["execution_mode"] == "HYBRID"


async def test_new_send_map_reduce_splits_only_the_failed_batch_and_reduces_recovery(
    tmp_path: Path,
) -> None:
    llm = SplittingMapReduceFixtureLlm()

    result = await run_consensus_fixture(tmp_path, llm)  # type: ignore[arg-type]

    TaskResultData.model_validate(result)
    assert len(llm.batch_payloads) >= 4


async def test_draft_review_downloads_and_parses_every_file_without_mocking(
    tmp_path: Path,
) -> None:
    bodies = {
        "/target.docx": build_docx(
            tmp_path / "target.docx", "融资租赁合同", "第一条 融资金额为1000万元。"
        ),
        "/template.docx": build_docx(
            tmp_path / "template.docx", "融资租赁合同", "第一条 融资金额为##{融资金额}万元。"
        ),
        "/reference.docx": build_docx(
            tmp_path / "reference.docx", "任意辅助资料", "仅参与本阶段解析。"
        ),
    }

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=bodies[request.url.path], request=request)

    settings = Settings(
        _env_file=None,
        TEMP_ROOT=str(tmp_path / "workspaces"),
        ALLOW_HTTP_DOWNLOADS=True,
        DOWNLOAD_HOST_ALLOWLIST="fixture-server",
    )
    executor = DraftReviewWorkflowExecutor(
        settings,
        downloader=SafeFileDownloadService(
            settings,
            transport=httpx.MockTransport(handler),
            resolver=resolver,
        ),
    )
    updates: list[tuple[TaskStage, int]] = []

    async def progress(stage: TaskStage, value: int, message: str) -> None:
        updates.append((stage, value))

    output = await executor.run(
        task_id="tsk_draft_parse",
        task_type=TaskType.DRAFT_REVIEW,
        files=[
            {
                "file_id": "fil_target",
                "role": "TARGET",
                "file_name": "target.docx",
                "url": "http://fixture-server/target.docx?token=secret",
                "safe_url": "http://fixture-server/target.docx",
            },
            {
                "file_id": "fil_template",
                "role": "TEMPLATE",
                "file_name": "template.docx",
                "url": "http://fixture-server/template.docx?token=secret",
                "safe_url": "http://fixture-server/template.docx",
            },
            {
                "file_id": "fil_reference",
                "role": "REFERENCE",
                "file_name": "reference.docx",
                "url": "http://fixture-server/reference.docx?token=secret",
                "safe_url": "http://fixture-server/reference.docx",
            },
        ],
        options={},
        progress_callback=progress,
    )

    result = output.result
    TaskResultData.model_validate(result)
    assert result["mock"] is False
    assert result["metadata"]["execution_mode"] == "RULE_BASED"
    assert result["schema_version"] == "2.1"
    assert result["metadata"]["workflow_version"] == "0.7.0"
    assert result["metadata"]["rules_version"] == "0.6.0"
    assert result["metadata"]["primary_model"] is None
    assert result["conclusion"] == "PASS"
    assert result["diff_items"] == []
    assert result["rule_checks"] == []
    assert result["metadata"]["template_diagnostics"]["filtered_diff_count"] == 1
    assert len({item["code"] for item in result["warnings"]}) == len(result["warnings"])
    assert len(result["files"]) == 3
    assert all(item["parser_name"] == "python-docx" for item in result["files"])
    assert all(item["document_profile"]["document_kind"] == "UNKNOWN" for item in result["files"])
    assert all(item["content_structure"]["block_count"] == 2 for item in result["files"])
    assert result["files"][0]["content_structure"]["sample_locations"]
    assert result["warnings"][-1]["code"] == "DRAFT_REVIEW_RULE_BASED_LIMITATION"
    assert "token=secret" not in str(result)
    assert updates[-1][0] == TaskStage.PERSISTING_RESULT
    assert any(stage == TaskStage.TEMPLATE_COMPARE for stage, _value in updates)
    assert not any((tmp_path / "workspaces").iterdir())
    assert len(output.file_metadata) == 3


@pytest.mark.asyncio
async def test_invalid_json_splits_only_the_failed_extraction_batch(
    tmp_path: Path,
) -> None:
    llm = SplitAfterJsonFailureLlm()
    result = await run_consensus_fixture(tmp_path, llm)

    assert result["metadata"]["execution_mode"] == "HYBRID"
    assert [len(payload["blocks"]) for payload in llm.payloads[:3]] == [2, 1, 1]
    assert len(llm.payloads) == 4


@pytest.mark.asyncio
async def test_invalid_json_stops_at_singleton_without_infinite_splitting(
    tmp_path: Path,
) -> None:
    llm = AlwaysJsonFailureLlm()

    with pytest.raises(WorkflowError) as error:
        await run_consensus_fixture(tmp_path, llm)

    assert error.value.code == "DYNAMIC_CHECK_INCOMPLETE"
    assert llm.calls == 2


@pytest.mark.parametrize(
    "settings_overrides",
    [
        {"LLM_EXTRACTION_MAX_REQUESTS_PER_DOCUMENT": 2},
        {"LLM_EXTRACTION_MAX_SPLIT_DEPTH": 0},
    ],
)
@pytest.mark.asyncio
async def test_invalid_json_respects_extraction_budget_and_depth(
    tmp_path: Path, settings_overrides: dict
) -> None:
    llm = AlwaysJsonFailureLlm()

    with pytest.raises(WorkflowError) as error:
        await run_consensus_fixture(
            tmp_path,
            llm,
            settings_overrides=settings_overrides,
        )

    assert error.value.code == "DYNAMIC_CHECK_INCOMPLETE"
    assert llm.calls == 1


@pytest.mark.asyncio
async def test_non_json_response_error_does_not_trigger_batch_split(
    tmp_path: Path,
) -> None:
    llm = EnvelopeFailureLlm()

    with pytest.raises(WorkflowError) as error:
        await run_consensus_fixture(tmp_path, llm)

    assert error.value.code == "DYNAMIC_CHECK_INCOMPLETE"
    assert llm.calls == 1


async def test_draft_review_uses_evidenced_llm_results_without_losing_rule_results(
    tmp_path: Path,
) -> None:
    bodies = {
        "/target.docx": build_docx(tmp_path / "target.docx", "合同", "固定条款"),
        "/template.docx": build_docx(tmp_path / "template.docx", "合同", "固定条款"),
        "/reference.docx": build_docx(tmp_path / "reference.docx", "合同", "辅助资料"),
    }

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=bodies[request.url.path], request=request)

    configured = Settings(
        _env_file=None,
        TEMP_ROOT=str(tmp_path / "workspaces"),
        ALLOW_HTTP_DOWNLOADS=True,
        DOWNLOAD_HOST_ALLOWLIST="fixture-server",
        LLM_ENABLED=True,
        LLM_BASE_URL="https://llm.invalid",
        LLM_API_KEY="unused",
    )
    executor = DraftReviewWorkflowExecutor(
        configured,
        downloader=SafeFileDownloadService(
            configured, transport=httpx.MockTransport(handler), resolver=resolver
        ),
        llm=EvidencedLlm(),
    )

    async def progress(stage: TaskStage, value: int, message: str) -> None:
        return None

    output = await executor.run(
        task_id="tsk_hybrid",
        task_type=TaskType.DRAFT_REVIEW,
        files=[
            {
                "file_id": "fil_target",
                "role": "TARGET",
                "file_name": "target.docx",
                "url": "http://fixture-server/target.docx",
                "safe_url": "http://fixture-server/target.docx",
            },
            {
                "file_id": "fil_template",
                "role": "TEMPLATE",
                "file_name": "template.docx",
                "url": "http://fixture-server/template.docx",
                "safe_url": "http://fixture-server/template.docx",
            },
            {
                "file_id": "fil_reference",
                "role": "REFERENCE",
                "file_name": "reference.docx",
                "url": "http://fixture-server/reference.docx",
                "safe_url": "http://fixture-server/reference.docx",
            },
        ],
        options={},
        progress_callback=progress,
    )

    result = output.result
    assert result["metadata"]["execution_mode"] == "HYBRID"
    TaskResultData.model_validate(result)
    assert result["metadata"]["primary_model"] == "test-model-v1"
    assert len(result["metadata"]["model_runs"]) == 2
    assert result["fact_matrix"][0]["status"] == "CONSISTENT"
    assert any(item["module_code"] == "FACT_CONSISTENCY" for item in result["passed_checks"])


async def test_dynamic_extraction_failure_logs_safe_category_only(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    events: list[tuple[str, dict]] = []

    def record_error(event: str, **kwargs: object) -> None:
        events.append((event, kwargs))

    monkeypatch.setattr(logger, "error", record_error)

    with pytest.raises(WorkflowError) as error:
        await run_consensus_fixture(tmp_path, ClassifiedFailureLlm())

    assert error.value.code == "DYNAMIC_CHECK_INCOMPLETE"
    assert error.value.details is None
    assert _dynamic_failure_code(
        LlmClientError("LLM_TIMEOUT", "upstream failed")
    ) == "LLM_UPSTREAM_FAILED"
    assert len(events) == 1
    event, fields = events[0]
    assert event == "draft_review_dynamic_check_failed"
    assert fields == {
        "task_id": "tsk_consensus_fixture",
        "stage": "FACT_EXTRACTION",
        "document_role": "TARGET",
        "chunk_index": 1,
        "error_category": "FACT_LOCATION_NOT_FOUND",
        "error_code": "LLM_EXTRACTION_EVIDENCE_INVALID",
        "affected_count": 1,
        "failure_counts": {"FACT_LOCATION_NOT_FOUND": 1},
        "split_depth": 0,
        "request_attempts": 1,
        "structure_retries": 0,
        "split_count": 0,
        "max_payload_chars": 757,
        "numeric_candidate_total": 1,
    }


@pytest.mark.parametrize(
    ("error_code", "expected_category"),
    [
        ("LLM_INVALID_JSON", "LLM_RESPONSE_JSON_INVALID"),
        ("LLM_RESPONSE_INVALID", "LLM_RESPONSE_ENVELOPE_INVALID"),
        ("LLM_SCHEMA_INVALID", "LLM_RESPONSE_SCHEMA_INVALID"),
    ],
)
def test_dynamic_failure_code_separates_response_failure_types(
    error_code: str, expected_category: str
) -> None:
    assert _dynamic_failure_code(LlmClientError(error_code, "safe failure")) == expected_category


def test_dynamic_failure_code_preserves_unknown_llm_error_code() -> None:
    assert _dynamic_failure_code(
        LlmClientError("LLM_SEMANTIC_PLAN_INVALID", "safe failure")
    ) == "LLM_SEMANTIC_PLAN_INVALID"


async def test_dynamic_schema_failure_logs_safe_validation_summary(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    events: list[tuple[str, dict]] = []

    def record_error(event: str, **kwargs: object) -> None:
        events.append((event, kwargs))

    monkeypatch.setattr(logger, "error", record_error)

    with pytest.raises(WorkflowError) as error:
        await run_consensus_fixture(tmp_path, SchemaFailureLlm())

    assert error.value.code == "DYNAMIC_CHECK_INCOMPLETE"
    assert error.value.details is None
    assert len(events) == 1
    fields = events[0][1]
    assert fields["error_category"] == "LLM_RESPONSE_SCHEMA_INVALID"
    assert fields["error_code"] == "LLM_SCHEMA_INVALID"
    assert fields["validation_summary_status"] == "PRESENT"
    assert fields["validation_summary"] == [
        {
            "path": "facts.*.value_type",
            "error_type": "literal_error",
            "count": 1,
        }
    ]


async def test_dynamic_schema_failure_without_summary_logs_missing_status(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    events: list[tuple[str, dict]] = []

    def record_error(event: str, **kwargs: object) -> None:
        events.append((event, kwargs))

    monkeypatch.setattr(logger, "error", record_error)

    with pytest.raises(WorkflowError) as error:
        await run_consensus_fixture(tmp_path, SchemaFailureWithoutSummaryLlm())

    assert error.value.code == "DYNAMIC_CHECK_INCOMPLETE"
    fields = events[0][1]
    assert fields["error_code"] == "LLM_SCHEMA_INVALID"
    assert fields["validation_summary_status"] == "MISSING"
    assert "validation_summary" not in fields


async def test_dynamic_raw_validation_error_logs_safe_summary(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    events: list[tuple[str, dict]] = []

    def record_error(event: str, **kwargs: object) -> None:
        events.append((event, kwargs))

    monkeypatch.setattr(logger, "error", record_error)

    with pytest.raises(WorkflowError) as error:
        await run_consensus_fixture(tmp_path, RawSchemaFailureLlm())

    assert error.value.code == "DYNAMIC_CHECK_INCOMPLETE"
    assert error.value.details is None
    fields = events[0][1]
    assert fields["validation_summary_status"] == "PRESENT"
    assert fields["validation_summary"] == [
        {
            "path": "facts.*.value_type",
            "error_type": "literal_error",
            "count": 1,
        }
    ]
    safe_text = str(fields["validation_summary"])
    assert "NOT_A_VALUE" not in safe_text
    assert "msg" not in safe_text
    assert "input" not in safe_text
    assert "ctx" not in safe_text


@pytest.mark.parametrize(
    "kwargs",
    [
        {"review_failure": True},
        {"extraction_model": "same-model", "review_model": "same-model"},
        {"extraction_confidence": 0.84},
        {"review_confidence": 0.84},
        {"evidence_complete": False},
    ],
)
async def test_consensus_gate_never_auto_accepts_unsafe_fact_result(
    tmp_path: Path, kwargs: dict
) -> None:
    with pytest.raises(WorkflowError, match="动态事实检查"):
        await run_consensus_fixture(tmp_path, ConsensusFixtureLlm(**kwargs))


async def test_invalid_advice_risk_id_falls_back_without_changing_result(
    tmp_path: Path,
) -> None:
    llm = ConsensusFixtureLlm(invalid_advice_evidence=True)
    result = await run_consensus_fixture(
        tmp_path,
        llm,
    )

    assert llm.advice_calls == 1
    assert result["conclusion"] == "PASS"
    assert result["review_items"] == []
    assert result["summary"]["statistics"]["review_count"] == 0
    assert any(
        warning["code"] == "LLM_ADVICE_UNAVAILABLE"
        for warning in result["warnings"]
    )


async def test_independent_review_receives_source_blocks(tmp_path: Path) -> None:
    llm = ConsensusFixtureLlm()

    await run_consensus_fixture(tmp_path, llm)

    assert llm.review_payloads
    assert all(payload["blocks"] for payload in llm.review_payloads)
    assert all(
        "text" in block and "location" in block
        for payload in llm.review_payloads
        for block in payload["blocks"]
    )
    assert all(
        payload["review_requirements"]["required_decision_count"]
        == len(payload["facts"])
        for payload in llm.review_payloads
    )


async def test_semantic_plan_runs_after_mapping_and_is_programmatically_checked(
    tmp_path: Path,
) -> None:
    llm = SemanticPlanFixtureLlm()
    result = await run_consensus_fixture(tmp_path, llm)

    assert llm.plan_calls == 1
    assert result["metadata"]["semantic_concepts"]
    assert any(
        spec["validation_id"] == "amount_positive"
        for spec in result["metadata"]["validation_specs"]
    )
    assert any(run["purpose"] == "SEMANTIC_PLAN" for run in result["metadata"]["model_runs"])
    assert any(check["rule_id"] == "amount_positive" for check in result["rule_checks"])


async def test_expanded_table_continues_into_facts_and_advice(tmp_path: Path) -> None:
    llm = ConsensusFixtureLlm()
    result = await run_consensus_fixture(tmp_path, llm, expanded_table=True)

    TaskResultData.model_validate(result)
    assert any(
        item["diff_type"] == "TABLE_STRUCTURE_EXPANDED" for item in result["diff_items"]
    )
    assert any(run["purpose"] == "FACT_EXTRACTION" for run in result["metadata"]["model_runs"])
    assert any(
        warning["code"] == "TEMPLATE_TABLE_STRUCTURE_EXPANDED"
        for warning in result["warnings"]
    )
    structure_diff_ids = {
        item["diff_id"]
        for item in result["diff_items"]
        if item["diff_type"] == "TABLE_STRUCTURE_EXPANDED"
    }
    structure_risks = [
        risk
        for risk in result["risk_items"]
        if structure_diff_ids.intersection(risk["related_diff_ids"])
    ]
    assert structure_risks
    assert all(risk.get("analysis_advice") for risk in structure_risks)


async def test_same_model_diagnostic_never_claims_independent_consensus_or_pass(
    tmp_path: Path,
) -> None:
    result = await run_consensus_fixture(
        tmp_path,
        ConsensusFixtureLlm(extraction_model="deepseek", review_model="deepseek"),
        same_model_diagnostic=True,
    )

    TaskResultData.model_validate(result)
    assert result["conclusion"] == "REVIEW_REQUIRED"
    assert result["metadata"]["execution_mode"] == "HYBRID_DIAGNOSTIC"
    assert result["metadata"]["independent_review"] is False
    assert result["metadata"]["review_mode"] == "SAME_MODEL_DIAGNOSTIC"
    assert result["review_items"]
    assert not any(
        item["module_code"] == "FACT_CONSISTENCY" for item in result["passed_checks"]
    )
    assert any(
        warning["code"] == "LLM_SAME_MODEL_DIAGNOSTIC"
        for warning in result["warnings"]
    )


async def test_same_model_semantic_plan_stays_out_of_public_formal_checks(
    tmp_path: Path,
) -> None:
    result = await run_consensus_fixture(
        tmp_path,
        SemanticPlanFixtureLlm(),
        same_model_diagnostic=True,
    )

    assert result["metadata"]["semantic_concepts"] == []
    assert result["metadata"]["validation_specs"] == []
    assert result["rule_checks"] == []
    assert not any(
        item.get("module_code") in {"NUMERIC_CONSISTENCY", "FACT_CONSISTENCY"}
        for item in result["risk_items"]
    )


async def test_same_model_diagnostic_preserves_deterministic_text_differences(
    tmp_path: Path,
) -> None:
    result = await run_consensus_fixture(
        tmp_path,
        ConsensusFixtureLlm(extraction_model="deepseek", review_model="deepseek"),
        same_model_diagnostic=True,
        target_body="固定条款乙",
        template_body="固定条款甲",
    )

    TaskResultData.model_validate(result)
    assert result["conclusion"] == "RISK_FOUND"
    assert result["diff_items"]
    assert any(item["module_code"] == "TEMPLATE_INTEGRITY" for item in result["risk_items"])


async def test_numeric_consistency_false_skips_dynamic_rules(tmp_path: Path) -> None:
    result = await run_consensus_fixture(
        tmp_path,
        ConsensusFixtureLlm(include_numeric_rule=True),
        options={"check_numeric_consistency": False},
    )

    assert not any(
        item.get("module_code") == "NUMERIC_CONSISTENCY" for item in result["risk_items"]
    )
    assert not any(
        item.get("module_code") == "NUMERIC_CONSISTENCY" for item in result["review_items"]
    )
    assert not any(
        item.get("module_code") == "NUMERIC_CONSISTENCY" for item in result["passed_checks"]
    )
    assert not any(
        check["rule_id"] == "amount_matches_literal" for check in result["rule_checks"]
    )


async def test_numeric_rule_requires_exact_primary_and_reviewer_consensus(
    tmp_path: Path,
) -> None:
    with pytest.raises(WorkflowError, match="动态数值检查"):
        await run_consensus_fixture(
            tmp_path,
            ConsensusFixtureLlm(include_numeric_rule=True, review_rule_mismatch=True),
        )
