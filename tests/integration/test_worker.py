import asyncio
from datetime import UTC, datetime, timedelta

import httpx
from docx import Document
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy import select

from app.core.config import get_settings
from app.core.enums import TaskStage, TaskStatus
from app.core.errors import WorkflowError
from app.db.models import CheckTask, TaskFile, TaskResult
from app.db.repositories.task_repository import TaskRepository
from app.db.session import SessionFactory
from app.main import app
from app.schemas.results import TaskResultData
from app.services.downloader import SafeFileDownloadService
from app.worker.runner import WorkerRunner
from app.workflows.draft_review import DraftReviewWorkflowExecutor
from app.workflows.final_compare import FinalCompareWorkflowExecutor
from app.workflows.mock_graphs import MockWorkflowExecutor
from tests.integration.helpers import DRAFT_PAYLOAD, FINAL_PAYLOAD


async def create(path: str, payload: dict) -> str:
    async with httpx.AsyncClient(
        transport=httpx.ASGITransport(app=app), base_url="http://test"
    ) as client:
        response = await client.post(path, json=payload)
        assert response.status_code == 202, response.text
        return response.json()["data"]["task_id"]


async def test_worker_claims_with_skip_locked_and_does_not_duplicate() -> None:
    task_id = await create("/api/v1/final-comparisons", FINAL_PAYLOAD)
    repo = TaskRepository()

    async def claim(worker: str):
        async with SessionFactory() as session, session.begin():
            task = await repo.claim_next(session, worker)
            return task.id if task else None

    first, second = await asyncio.gather(claim("worker-a"), claim("worker-b"))
    assert sorted(value for value in (first, second) if value) == [task_id]


async def test_full_worker_success_and_atomic_result() -> None:
    task_id = await create("/api/v1/draft-reviews", DRAFT_PAYLOAD)
    settings = get_settings()
    runner = WorkerRunner(settings, workflow=MockWorkflowExecutor(settings))
    assert await runner.run_once() is True

    async with SessionFactory() as session:
        task = (
            await session.execute(select(CheckTask).where(CheckTask.id == task_id))
        ).scalar_one()
        result = (
            await session.execute(select(TaskResult).where(TaskResult.task_id == task_id))
        ).scalar_one()
        files = (
            (await session.execute(select(TaskFile).where(TaskFile.task_id == task_id)))
            .scalars()
            .all()
        )
    assert task.status == TaskStatus.SUCCEEDED
    assert task.progress == 100 and task.finished_at is not None
    assert result.result["mock"] is True
    TaskResultData.model_validate(result.result)
    assert all("?" not in file.safe_url for file in files)
    assert "target-secret" not in str(task.input_snapshot)


async def test_worker_failure_then_retry_creates_new_task() -> None:
    task_id = await create("/api/v1/draft-reviews", DRAFT_PAYLOAD)
    settings = get_settings()
    runner = WorkerRunner(
        settings,
        workflow=MockWorkflowExecutor(settings, fail_stage=TaskStage.PARSING),
    )
    await runner.run_once()
    async with httpx.AsyncClient(
        transport=httpx.ASGITransport(app=app), base_url="http://test"
    ) as client:
        failed = await client.get(f"/api/v1/tasks/{task_id}")
        assert failed.json()["data"]["status"] == "FAILED"
        response = await client.post(f"/api/v1/tasks/{task_id}/retry")
        assert response.status_code == 202, response.text
        retried = response.json()["data"]
        assert retried["task_id"] != task_id
        assert retried["source_task_id"] == task_id
        assert retried["status"] == "PENDING"


async def test_worker_persists_and_returns_safe_workflow_error_details(capsys) -> None:
    class SafeFailureWorkflow:
        async def run(self, **kwargs):
            raise WorkflowError(
                "OCR_SERVICE_UNAVAILABLE",
                "OCR 服务暂时不可用",
                details={
                    "component": "EXTERNAL_DOCUMENT_PARSER",
                    "failure_kind": "UPSTREAM_502",
                    "attempts": 1,
                    "elapsed_ms": 25,
                },
            )

    task_id = await create("/api/v1/final-comparisons", FINAL_PAYLOAD)
    runner = WorkerRunner(get_settings(), workflow=SafeFailureWorkflow())
    assert await runner.run_once() is True

    async with httpx.AsyncClient(
        transport=httpx.ASGITransport(app=app), base_url="http://test"
    ) as client:
        detail = (await client.get(f"/api/v1/tasks/{task_id}")).json()["data"]

    assert detail["error"] == {
        "code": "OCR_SERVICE_UNAVAILABLE",
        "message": "OCR 服务暂时不可用",
        "details": {
            "component": "EXTERNAL_DOCUMENT_PARSER",
            "failure_kind": "UPSTREAM_502",
            "attempts": 1,
            "elapsed_ms": 25,
        },
    }
    output = capsys.readouterr().out
    assert "UPSTREAM_502" in output
    assert "OCR_BASE_URL" not in output


async def test_stale_task_requeues_then_fails_at_max_attempts() -> None:
    task_id = await create("/api/v1/final-comparisons", FINAL_PAYLOAD)
    old = datetime.now(UTC) - timedelta(hours=1)
    async with SessionFactory() as session, session.begin():
        task = (
            await session.execute(select(CheckTask).where(CheckTask.id == task_id))
        ).scalar_one()
        task.status = TaskStatus.RUNNING
        task.stage = TaskStage.PARSING
        task.worker_id = "dead-worker"
        task.heartbeat_at = old
        task.attempt_count = 1
        task.max_attempts = 2
    runner = WorkerRunner(get_settings())
    requeued, failed = await runner.recover_stale()
    assert requeued == [task_id] and failed == []

    async with SessionFactory() as session, session.begin():
        task = (
            await session.execute(select(CheckTask).where(CheckTask.id == task_id))
        ).scalar_one()
        task.status = TaskStatus.RUNNING
        task.stage = TaskStage.PARSING
        task.worker_id = "dead-worker-2"
        task.heartbeat_at = old
        task.attempt_count = 2
    requeued, failed = await runner.recover_stale()
    assert requeued == [] and failed == [task_id]
    async with SessionFactory() as session:
        task = (
            await session.execute(select(CheckTask).where(CheckTask.id == task_id))
        ).scalar_one()
    assert task.status == TaskStatus.FAILED and task.error_code == "WORKER_LOST"


async def test_real_final_compare_worker_persists_result_and_file_metadata(tmp_path) -> None:
    baseline_path = tmp_path / "baseline.docx"
    target_path = tmp_path / "target.docx"
    for path, amount in ((baseline_path, "100"), (target_path, "120")):
        document = Document()
        document.add_paragraph(f"合同金额为{amount}万元，期限为24个月。")
        document.save(path)
    bodies = {
        "/baseline.docx": baseline_path.read_bytes(),
        "/target.docx": target_path.read_bytes(),
    }

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=bodies[request.url.path], request=request)

    async def resolver(host: str, port: int) -> list[str]:
        return ["127.0.0.1"]

    payload = {
        "client_reference_id": "integration-real-final",
        "baseline_file": {
            "url": "http://fixture-server/baseline.docx?token=secret",
            "file_name": "baseline.docx",
        },
        "target_file": {
            "url": "http://fixture-server/target.docx?token=secret",
            "file_name": "target.docx",
        },
    }
    task_id = await create("/api/v1/final-comparisons", payload)
    base_settings = get_settings()
    real_settings = base_settings.model_copy(
        update={
            "TEMP_ROOT": str(tmp_path / "work"),
            "ALLOW_HTTP_DOWNLOADS": True,
            "DOWNLOAD_HOST_ALLOWLIST": "fixture-server",
            "OCR_ENABLED": False,
        }
    )
    downloader = SafeFileDownloadService(
        real_settings,
        transport=httpx.MockTransport(handler),
        resolver=resolver,
    )
    runner = WorkerRunner(
        real_settings,
        workflow=FinalCompareWorkflowExecutor(real_settings, downloader=downloader),
    )
    assert await runner.run_once() is True

    async with SessionFactory() as session:
        task = (
            await session.execute(select(CheckTask).where(CheckTask.id == task_id))
        ).scalar_one()
        stored = (
            await session.execute(select(TaskResult).where(TaskResult.task_id == task_id))
        ).scalar_one()
        files = (
            (await session.execute(select(TaskFile).where(TaskFile.task_id == task_id)))
            .scalars()
            .all()
        )
    assert task.status == TaskStatus.SUCCEEDED
    assert stored.result["mock"] is False
    assert stored.result["metadata"]["execution_mode"] == "RULE_BASED"
    assert stored.model_name is None
    assert all(file.sha256 and file.parser_name == "python-docx" for file in files)
    assert "token=secret" not in str(stored.result)


async def test_real_draft_review_worker_parses_all_files_and_persists_metadata(tmp_path) -> None:
    bodies: dict[str, bytes] = {}
    for name, title in (
        ("target.docx", "目标合同"),
        ("template.docx", "合同模板"),
        ("reference.docx", "任意辅助资料"),
    ):
        path = tmp_path / name
        document = Document()
        document.add_heading(title, level=1)
        document.add_paragraph("用于集成测试的合成正文。")
        document.save(path)
        bodies[f"/{name}"] = path.read_bytes()

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=bodies[request.url.path], request=request)

    async def resolver(host: str, port: int) -> list[str]:
        return ["127.0.0.1"]

    payload = {
        "client_reference_id": "integration-real-draft",
        "target_file": {
            "url": "http://fixture-server/target.docx?token=secret",
            "file_name": "target.docx",
        },
        "template_file": {
            "url": "http://fixture-server/template.docx?token=secret",
            "file_name": "template.docx",
        },
        "reference_files": [
            {
                "url": "http://fixture-server/reference.docx?token=secret",
                "file_name": "reference.docx",
            }
        ],
    }
    task_id = await create("/api/v1/draft-reviews", payload)
    real_settings = get_settings().model_copy(
        update={
            "TEMP_ROOT": str(tmp_path / "work-draft"),
            "ALLOW_HTTP_DOWNLOADS": True,
            "DOWNLOAD_HOST_ALLOWLIST": "fixture-server",
            "OCR_ENABLED": False,
        }
    )
    downloader = SafeFileDownloadService(
        real_settings,
        transport=httpx.MockTransport(handler),
        resolver=resolver,
    )
    runner = WorkerRunner(
        real_settings,
        workflow=DraftReviewWorkflowExecutor(real_settings, downloader=downloader),
    )
    assert await runner.run_once() is True

    async with SessionFactory() as session:
        task = (
            await session.execute(select(CheckTask).where(CheckTask.id == task_id))
        ).scalar_one()
        stored = (
            await session.execute(select(TaskResult).where(TaskResult.task_id == task_id))
        ).scalar_one()
        files = (
            (await session.execute(select(TaskFile).where(TaskFile.task_id == task_id)))
            .scalars()
            .all()
        )

    assert task.status == TaskStatus.SUCCEEDED
    assert stored.result["mock"] is False
    assert stored.result["metadata"]["execution_mode"] == "PARSER_ONLY"
    assert stored.result["conclusion"] == "REVIEW_REQUIRED"
    assert len(stored.result["files"]) == 3
    assert stored.model_name is None
    assert all(file.sha256 and file.parser_name == "python-docx" for file in files)
    assert all(file.reference_type is None for file in files)
    assert "token=secret" not in str(stored.result)
    assert not any((tmp_path / "work-draft").iterdir())


async def test_formal_pdf_without_external_parser_fails_safely_and_cleans_workspace(
    tmp_path,
) -> None:
    empty_pdf = tmp_path / "empty.pdf"
    canvas = Canvas(str(empty_pdf))
    canvas.showPage()
    canvas.save()

    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=empty_pdf.read_bytes(), request=request)

    async def resolver(host: str, port: int) -> list[str]:
        return ["127.0.0.1"]

    payload = {
        "client_reference_id": "integration-empty-pdf",
        "baseline_file": {"url": "http://fixture-server/base.pdf", "file_name": "base.pdf"},
        "target_file": {"url": "http://fixture-server/target.pdf", "file_name": "target.pdf"},
    }
    task_id = await create("/api/v1/final-comparisons", payload)
    work_root = tmp_path / "work"
    real_settings = get_settings().model_copy(
        update={
            "TEMP_ROOT": str(work_root),
            "ALLOW_HTTP_DOWNLOADS": True,
            "DOWNLOAD_HOST_ALLOWLIST": "fixture-server",
        }
    )
    downloader = SafeFileDownloadService(
        real_settings,
        transport=httpx.MockTransport(handler),
        resolver=resolver,
    )
    runner = WorkerRunner(
        real_settings,
        workflow=FinalCompareWorkflowExecutor(real_settings, downloader=downloader),
    )
    assert await runner.run_once() is True
    async with SessionFactory() as session:
        task = (
            await session.execute(select(CheckTask).where(CheckTask.id == task_id))
        ).scalar_one()
    assert task.status == TaskStatus.FAILED
    assert task.error_code == "OCR_NOT_CONFIGURED"
    assert work_root.exists() and not any(work_root.iterdir())
